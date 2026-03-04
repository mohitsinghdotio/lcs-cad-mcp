"""Report generation service — PDF, DOCX, and JSON reports from ScrutinyReport."""
from __future__ import annotations

import json
import logging
from pathlib import Path
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from lcs_cad_mcp.rule_engine.models import ScrutinyReport

logger = logging.getLogger(__name__)


class ReportGenerationError(Exception):
    """Raised when report generation fails."""
    def __init__(self, code: str, message: str) -> None:
        super().__init__(message)
        self.code = code
        self.message = message


class ReportService:
    """Generates PDF, DOCX, and JSON scrutiny reports from a ScrutinyReport."""

    def assemble_data(self, report: ScrutinyReport, run_id: str | None = None) -> dict:
        """Assemble a report data dict from a ScrutinyReport."""
        return {
            "run_id": run_id or report.run_id,
            "drawing_path": report.drawing_path,
            "authority": report.authority,
            "overall_pass": report.overall_pass,
            "total_rules": report.total_rules,
            "passed_rules": report.passed_rules,
            "failed_rules": report.failed_rules,
            "results": [
                {
                    "rule_id": r.rule_id,
                    "rule_name": r.rule_name,
                    "status": r.status,
                    "computed_value": r.computed_value,
                    "permissible_value": r.permissible_value,
                    "deviation": r.deviation,
                    "suggested_action": r.suggested_action,
                }
                for r in report.results
            ],
        }

    def generate_json(self, report: ScrutinyReport, output_path: str,
                      run_id: str | None = None) -> str:
        """Write a JSON scrutiny report to output_path."""
        data = self.assemble_data(report, run_id)
        path = Path(output_path)
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(json.dumps(data, indent=2), encoding="utf-8")
        logger.info("JSON report written to %s", output_path)
        return output_path

    def generate_pdf(self, report: ScrutinyReport, output_path: str,
                     run_id: str | None = None) -> str:
        """Write a PDF scrutiny report to output_path using reportlab."""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            from reportlab.lib import colors
        except ImportError:
            raise ReportGenerationError(
                code="REPORTLAB_NOT_AVAILABLE",
                message="reportlab is required for PDF generation. Install it with: pip install reportlab",
            )

        data = self.assemble_data(report, run_id)
        path = Path(output_path)
        path.parent.mkdir(parents=True, exist_ok=True)

        doc = SimpleDocTemplate(str(path), pagesize=A4)
        styles = getSampleStyleSheet()
        story = []

        # Title
        story.append(Paragraph(f"AutoDCR Scrutiny Report", styles["Title"]))
        story.append(Spacer(1, 12))
        story.append(Paragraph(f"Authority: {data['authority']}", styles["Normal"]))
        story.append(Paragraph(f"Run ID: {data['run_id']}", styles["Normal"]))
        overall = "PASS" if data["overall_pass"] else "FAIL"
        story.append(Paragraph(f"Overall Result: {overall}", styles["Heading2"]))
        story.append(Spacer(1, 12))

        # Rules table
        table_data = [["Rule ID", "Rule Name", "Status", "Computed", "Allowed", "Deviation"]]
        for r in data["results"]:
            table_data.append([
                r["rule_id"], r["rule_name"], r["status"].upper(),
                str(round(r["computed_value"], 3)),
                str(round(r["permissible_value"], 3)),
                str(round(r["deviation"], 3)),
            ])

        table = Table(table_data, repeatRows=1)
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
        ]))
        story.append(table)

        doc.build(story)
        logger.info("PDF report written to %s", output_path)
        return output_path

    def generate_docx(self, report: ScrutinyReport, output_path: str,
                      run_id: str | None = None) -> str:
        """Write a DOCX scrutiny report to output_path using python-docx."""
        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError:
            raise ReportGenerationError(
                code="DOCX_NOT_AVAILABLE",
                message="python-docx is required for DOCX generation. Install it with: pip install python-docx",
            )

        data = self.assemble_data(report, run_id)
        path = Path(output_path)
        path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        doc.add_heading("AutoDCR Scrutiny Report", 0)
        doc.add_paragraph(f"Authority: {data['authority']}")
        doc.add_paragraph(f"Run ID: {data['run_id']}")
        overall = "PASS" if data["overall_pass"] else "FAIL"
        doc.add_heading(f"Overall Result: {overall}", level=1)

        # Rules table
        table = doc.add_table(rows=1, cols=6)
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(["Rule ID", "Rule Name", "Status", "Computed", "Allowed", "Deviation"]):
            hdr_cells[i].text = header

        for r in data["results"]:
            row_cells = table.add_row().cells
            row_cells[0].text = r["rule_id"]
            row_cells[1].text = r["rule_name"]
            row_cells[2].text = r["status"].upper()
            row_cells[3].text = str(round(r["computed_value"], 3))
            row_cells[4].text = str(round(r["permissible_value"], 3))
            row_cells[5].text = str(round(r["deviation"], 3))

        doc.save(str(path))
        logger.info("DOCX report written to %s", output_path)
        return output_path
